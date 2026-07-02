"""Render structured resume data into a professional, ATS-friendly .docx.

ATS-friendly: single column, real heading text, standard fonts, no text boxes.
For one-page templates we compress margins, spacing, and (adaptively) the font
size so dense content still fits a single page without dropping information.
"""

import io
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.resume_templates import effective_style

PAGE_WIDTH = 8.5  # US Letter inches


def _hex(c):
    return RGBColor.from_string(c)


def _add_md_runs(paragraph, text, size=None):
    """Add text to a paragraph, rendering **bold** spans as bold runs."""
    text = str(text or "")
    # Unbalanced markers -> render literally to avoid wrong toggling.
    segments = text.split("**") if text.count("**") % 2 == 0 else [text]
    for i, seg in enumerate(segments):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True
        if size:
            run.font.size = Pt(size)
    if not paragraph.runs:
        paragraph.add_run("")


def _add_bottom_border(paragraph, color):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    bdr.append(bottom)
    pPr.append(bdr)


def _estimate_lines(data, style):
    """Rough number of rendered text lines, used to pick a fitting font size."""
    usable_in = PAGE_WIDTH - 2 * style["margin"]
    # Conservative density (err toward over-counting so we shrink, never overflow).
    chars_per_line = max(36, int(usable_in * 11))
    lines = 4  # name + contact + padding
    order = style["section_order"]

    def text_lines(s):
        return max(1, math.ceil(len(str(s)) / chars_per_line))

    if style["include_summary"] and data.get("summary"):
        lines += 1 + text_lines(data["summary"])
    if "skills" in order and data.get("skills"):
        lines += 1 + text_lines(" • ".join(data["skills"]))
    if "experience" in order:
        for job in data.get("experience", []):
            lines += 2  # title line + location
            for b in job.get("bullets", []):
                if b:
                    lines += text_lines(b)
        if data.get("experience"):
            lines += 1  # heading
    if "projects" in order and data.get("projects"):
        lines += 1
        for p in data["projects"]:
            lines += 1 + text_lines(p.get("description", ""))
            lines += sum(1 for b in p.get("bullets", []) or [] if b)
    if "education" in order and data.get("education"):
        lines += 1 + len(data["education"])
    if "achievements" in order and data.get("achievements"):
        lines += 1 + sum(text_lines(a) for a in data["achievements"])
    if "certifications" in order and data.get("certifications"):
        lines += 1 + len(data["certifications"])
    return lines


def _fit(data, style):
    """Return (body_size, scale) adapted to keep one-page templates on one page."""
    body = style["body_size"]
    if not style.get("one_page"):
        return body, 1.0
    lines = _estimate_lines(data, style)
    budget = 46  # lines that comfortably fit one page at full size + 0.5" margins
    if lines <= budget:
        return body, 1.0
    # Step the font/spacing down progressively as content grows.
    for size, scale, cap in [(10.0, 0.85, 52), (9.5, 0.7, 60), (9.0, 0.55, 999)]:
        if lines <= cap:
            return size, scale
    return 8.5, 0.5


def render_docx(data: dict, template: str = "technical", fit: dict | None = None,
                layout: dict | None = None) -> bytes:
    style = effective_style(template, layout)
    if fit:  # use the PDF-measured fit so .docx and .pdf stay consistent
        body_size, scale, margin = fit["body"], fit["scale"], fit["margin"]
    else:
        body_size, scale = _fit(data, style)
        margin = style["margin"]
    doc = Document()

    sec = doc.sections[0]
    m = Inches(margin)
    sec.top_margin = sec.bottom_margin = m
    sec.left_margin = sec.right_margin = m
    right_tab = Inches(PAGE_WIDTH - 2 * margin)

    normal = doc.styles["Normal"]
    normal.font.name = style["font"]
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(2 * scale)

    def sp(before=0, after=2):
        return Pt(before * scale), Pt(after * scale)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before, p.paragraph_format.space_after = sp(7, 3)
        label = text.upper() if style["uppercase_headings"] else text
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(body_size + 1.5)
        r.font.name = style["font"]
        r.font.color.rgb = _hex(style["accent"])
        if style["heading_rule"]:
            _add_bottom_border(p, style["accent"])

    def title_dates(bold_text, rest, right):
        p = doc.add_paragraph()
        p.paragraph_format.space_before, p.paragraph_format.space_after = sp(2, 0)
        p.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(bold_text).bold = True
        if rest:
            p.add_run(rest)
        if right:
            rr = p.add_run(f"\t{right}")
            rr.italic = True
            rr.font.size = Pt(body_size - 1)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1 * scale)
        _add_md_runs(p, text)

    contact = data.get("contact", {}) or {}

    # Name + contact
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_after = Pt(1)
    nr = np.add_run(contact.get("full_name") or "Your Name")
    nr.bold = True
    nr.font.size = Pt(style["name_size"])
    nr.font.name = style["font"]
    nr.font.color.rgb = _hex(style["accent"])

    bits = [contact.get(k) for k in ("email", "phone", "location", "linkedin", "website")]
    line = "  |  ".join(b for b in bits if b)
    if line:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4 * scale)
        cp.add_run(line).font.size = Pt(body_size - 1)

    def render_section(key):
        if key == "summary" and style["include_summary"] and data.get("summary"):
            heading("Summary")
            _add_md_runs(doc.add_paragraph(), data["summary"])
        elif key == "skills" and data.get("skills"):
            heading("Skills")
            doc.add_paragraph(" • ".join(data["skills"]))
        elif key == "experience" and data.get("experience"):
            heading("Experience")
            for job in data["experience"]:
                company = job.get("company", "")
                dates = " – ".join(d for d in [job.get("start_date"), job.get("end_date")] if d)
                title_dates(job.get("title", ""), f"  —  {company}" if company else "", dates)
                if job.get("location"):
                    lp = doc.add_paragraph()
                    lp.paragraph_format.space_after = Pt(0)
                    lr = lp.add_run(job["location"])
                    lr.italic = True
                    lr.font.size = Pt(body_size - 1)
                for b in job.get("bullets", []):
                    if b:
                        bullet(b)
        elif key == "projects" and data.get("projects"):
            heading("Projects")
            for proj in data["projects"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1 * scale)
                p.add_run(proj.get("name", "")).bold = True
                if proj.get("description"):
                    p.add_run(f" — {proj['description']}")
                for b in proj.get("bullets", []) or []:
                    if b:
                        bullet(b)
        elif key == "education" and data.get("education"):
            heading("Education")
            for edu in data["education"]:
                inst = edu.get("institution", "")
                title_dates(edu.get("degree", ""), f"  —  {inst}" if inst else "",
                            edu.get("graduation_date", ""))
                if edu.get("details"):
                    dp = doc.add_paragraph()
                    dp.add_run(edu["details"]).font.size = Pt(body_size - 1)
        elif key == "achievements" and data.get("achievements"):
            heading("Achievements")
            for a in data["achievements"]:
                if a:
                    bullet(a)
        elif key == "certifications" and data.get("certifications"):
            heading("Certifications")
            for c in data["certifications"]:
                if c:
                    bullet(c)

    for key in style["section_order"]:
        render_section(key)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
