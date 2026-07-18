"""Extract raw text from an uploaded resume (PDF / DOCX / TXT)."""

import io

from docx import Document
from pypdf import PdfReader

SUPPORTED = (".pdf", ".docx", ".txt")


def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    if name.endswith(".txt"):
        return content.decode("utf-8", errors="ignore").strip()
    raise ValueError(
        f"Unsupported file type. Please upload one of: {', '.join(SUPPORTED)}"
    )
